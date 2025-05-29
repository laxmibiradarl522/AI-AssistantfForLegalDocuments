from flask import Flask, request, jsonify, render_template, session, send_file
import os
import openai
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import tempfile

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

app.secret_key = os.urandom(24)

openai.api_key=OPENAI_API_KEY

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    """Extracts text from a Word (.docx) file."""
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def summarize_text(text):
    """Summarizes the given text using OpenAI's GPT model."""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo", 
            messages=[{"role": "system", "content": "You are a helpful assistant that summarizes text."},
                      {"role": "user", "content": f"Summarize the following text:\n{text}"}],
            max_tokens=200,
            temperature=0.7
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"An error occurred: {str(e)}"

@app.route('/')
def index():
    """Renders the homepage."""
    return render_template('index.html')

@app.route('/summary')
def summary_page():
    """Renders the page for document summarization."""
    return render_template('summary.html')

@app.route('/chat')
def chat_page():
    """Renders the page for asking queries."""
    return render_template('chat.html')

@app.route('/document')
def document_page():
    """Renders the page for document generation"""
    return render_template('form.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handles file uploads and returns a summary of the uploaded document."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                return jsonify({'error': 'Unsupported file type'}), 400
            
            summary = summarize_text(text)
            
            os.remove(file_path)
            
            return jsonify({'summary': summary})
        
        except Exception as e:
            os.remove(file_path)  
            return jsonify({'error': f"An error occurred: {str(e)}"}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/ask_query', methods=['POST'])
def ask_query():
    """Handles user queries for the chatbot."""
    data = request.get_json()
    query = data.get('query')

    if not query:
        return jsonify({'error': 'No query provided'}), 400

    if 'conversation_history' not in session:
        session['conversation_history'] = []

    session['conversation_history'].append({"role": "user", "content": query})

    messages = [{"role": "system", "content": "You are a helpful assistant."}]
    messages.extend(session['conversation_history'])

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",  
            messages=messages,
            max_tokens=150,
            temperature=0.7
        )
        answer = response['choices'][0]['message']['content'].strip()

        session['conversation_history'].append({"role": "assistant", "content": answer})

        return jsonify({'answer': answer})

    except Exception as e:
        return jsonify({'error': f"An error occurred: {str(e)}"}), 500

@app.route('/generate', methods=['POST'])
def generate_document():
    user_input = request.form['user_input']

    prompt = f"Generate a legally valid document based on this user input:\n\n{user_input}\n\nThe output should be clear, formal, and suitable for legal use."

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # or "gpt-4" if you have access
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4
    )

    generated_text = response.choices[0].message.content.strip()

    # Create Word doc
    doc = Document()
    doc.add_heading('Generated Legal Document', level=1)
    doc.add_paragraph(generated_text)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)

    return send_file(temp_file.name, as_attachment=True, download_name="LegalDocument.docx")


if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)
